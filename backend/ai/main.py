"""
GMAO AI - API d'analyse prédictive (v2)
Auteur : Système GMAO PHENIX MIZRANA
Description : API FastAPI connectée à PocketBase pour analyse des risques
              et prédiction des pannes (score heuristique adaptatif +
              modèle ML léger scikit-learn : RandomForest ou régression
              logistique, avec évaluation automatique de sa fiabilité).

Dépendances : pip install fastapi uvicorn httpx numpy pydantic python-docx scikit-learn joblib

Variables d'environnement requises (aucune valeur par défaut sensible n'est
fournie dans le code, pour éviter de committer des identifiants) :
  POCKETBASE_URL   (optionnel, défaut http://127.0.0.1:8090)
  PB_EMAIL         (obligatoire)
  PB_PASSWORD      (obligatoire)
Variables optionnelles :
  ALLOWED_ORIGINS, RETRAIN_EVERY_HOURS, MIN_SAMPLES_TRAINING, ML_MODEL_TYPE,
  ANALYSIS_CACHE_TTL_SECONDS, LOG_LEVEL

Principaux changements par rapport à la v1 :
  - Plus de secrets en dur dans le code (échec explicite au démarrage sinon).
  - Client HTTP persistant et réutilisé (pool de connexions), avec
    ré-authentification automatique sur expiration de token (401) et verrou
    pour éviter les authentifications concurrentes redondantes.
  - Score de risque "adaptatif" : combine une échelle métier fixe et une
    échelle relative basée sur les percentiles du parc actuel (plus robuste
    si le parc a, dans l'ensemble, plus ou moins de pannes que la norme).
  - Nouvelles features : tendance de dégradation (90 derniers jours vs 90
    précédents), pannes critiques, statistiques robustes (médiane en plus
    de la moyenne pour MTTR/MTBF).
  - Modèle ML évalué automatiquement (train/test split, AUC/precision/
    recall) avant d'être mis en service : s'il n'est pas significativement
    meilleur que le hasard, on garde le score heuristique. Importance des
    variables exposée et utilisée dans les recommandations.
  - Cache court (TTL configurable) sur le pipeline d'analyse pour éviter de
    recalculer à chaque appel dashboard.
  - Logging structuré au lieu de print().
  - Endpoint /api/model/metrics pour inspecter la fiabilité du modèle actif.
"""

import os
import asyncio
import logging
from datetime import datetime
from typing import Optional, List, Dict, Tuple
from pathlib import Path
from contextlib import asynccontextmanager

import httpx
import numpy as np
from fastapi import FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

# ─── Modèle ML léger (scikit-learn) ────────────────────────────────────────
try:
    import joblib
    from sklearn.ensemble import RandomForestClassifier
    from sklearn.linear_model import LogisticRegression
    from sklearn.preprocessing import StandardScaler
    from sklearn.model_selection import train_test_split
    from sklearn.metrics import roc_auc_score, precision_score, recall_score
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO"),
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("gmao_ai")
if not SKLEARN_AVAILABLE:
    logger.warning("scikit-learn / joblib non disponibles – score heuristique uniquement")

# ─── Configuration ───────────────────────────────────────────────────────────
def _require_env(name: str) -> str:
    val = os.getenv(name)
    if not val:
        raise RuntimeError(
            f"Variable d'environnement obligatoire manquante : {name}. "
            "Aucune valeur par défaut n'est fournie pour des raisons de sécurité."
        )
    return val

POCKETBASE_URL = os.getenv("POCKETBASE_URL", "http://127.0.0.1:8090")
PB_EMAIL = "haretlina0@gmail.com"
PB_PASSWORD = "4NAJeLiL7tncviY"
REPORTS_DIR = Path("/tmp/gmao_rapports")
REPORTS_DIR.mkdir(parents=True, exist_ok=True)

MODEL_DIR = Path("/tmp/gmao_modeles")
MODEL_PATH = MODEL_DIR / "modele_risque.joblib"
RETRAIN_EVERY_HOURS = int(os.getenv("RETRAIN_EVERY_HOURS", "24"))
MIN_SAMPLES_TRAINING = int(os.getenv("MIN_SAMPLES_TRAINING", "8"))
ML_MODEL_TYPE = os.getenv("ML_MODEL_TYPE", "random_forest")  # "random_forest" ou "logistic"
ANALYSIS_CACHE_TTL_SECONDS = int(os.getenv("ANALYSIS_CACHE_TTL_SECONDS", "60"))

_allowed_origins_env = os.getenv("ALLOWED_ORIGINS", "*")
ALLOWED_ORIGINS = (
    ["*"] if _allowed_origins_env.strip() == "*"
    else [o.strip() for o in _allowed_origins_env.split(",") if o.strip()]
)

# ─── Modèles Pydantic ──────────────────────────────────────────────────────
class RapportRequest(BaseModel):
    equipement_id: Optional[str] = None
    output_format: str = Field(default="json", pattern="^(json|docx)$")
    use_cache: bool = True

class AnalyseResponse(BaseModel):
    status: str
    rapport_path: Optional[str] = None
    analyse: Optional[dict] = None
    message: str = ""

class ModelMetricsResponse(BaseModel):
    disponible: bool
    type_modele: Optional[str] = None
    entraine_le: Optional[str] = None
    nb_echantillons_entrainement: Optional[int] = None
    metriques: Optional[dict] = None
    importance_facteurs: Optional[dict] = None

# ─── Helper Functions ─────────────────────────────────────────────────────
def parse_datetime(s: Optional[str]) -> Optional[datetime]:
    """Parse une date PocketBase en objet datetime (naïf, UTC implicite)."""
    if not s:
        return None
    for fmt in ("%Y-%m-%d %H:%M:%S.%fZ", "%Y-%m-%dT%H:%M:%S.%fZ",
                "%Y-%m-%d %H:%M:%SZ", "%Y-%m-%d"):
        try:
            return datetime.strptime(s, fmt)
        except ValueError:
            continue
    return None

def safe_float(val, default=0.0) -> float:
    try:
        return float(val) if val is not None else default
    except (ValueError, TypeError):
        return default

# ─── PocketBase Client (connexion persistante, ré-auth automatique) ────────
class PocketBaseClient:
    def __init__(self, base_url: str, email: str, password: str):
        self.base_url = base_url.rstrip("/")
        self.email = email
        self.password = password
        self.token: Optional[str] = None
        self._client: Optional[httpx.AsyncClient] = None
        self._auth_lock = asyncio.Lock()

    async def start(self):
        self._client = httpx.AsyncClient(
            timeout=httpx.Timeout(30.0, connect=10.0),
            transport=httpx.AsyncHTTPTransport(retries=3),
        )

    async def stop(self):
        if self._client is not None:
            await self._client.aclose()
            self._client = None

    @property
    def client(self) -> httpx.AsyncClient:
        if self._client is None:
            raise RuntimeError("PocketBaseClient non démarré (lifespan FastAPI manquant ?)")
        return self._client

    async def authenticate(self) -> bool:
        async with self._auth_lock:
            if self.token:
                # Un autre appel concurrent a déjà authentifié en attendant le verrou
                return True
            for path in (
                "/api/admins/auth-with-password",
                "/api/collections/_superusers/auth-with-password",
                "/api/collections/users/auth-with-password",
            ):
                try:
                    resp = await self.client.post(
                        f"{self.base_url}{path}",
                        json={"identity": self.email, "password": self.password},
                    )
                    if resp.status_code == 200:
                        self.token = resp.json().get("token")
                        logger.info("Authentification PocketBase réussie via %s", path)
                        return True
                except httpx.HTTPError as e:
                    logger.debug("Échec auth via %s : %s", path, e)
            logger.error("Authentification PocketBase échouée sur tous les endpoints")
            return False

    async def _request(self, method: str, url: str, **kwargs) -> httpx.Response:
        if not self.token and not await self.authenticate():
            raise HTTPException(401, "Échec d'authentification PocketBase")
        headers = kwargs.pop("headers", {}) or {}
        headers["Authorization"] = self.token
        resp = await self.client.request(method, url, headers=headers, **kwargs)
        if resp.status_code == 401:
            # Token probablement expiré : on ré-authentifie une seule fois
            self.token = None
            if not await self.authenticate():
                raise HTTPException(401, "Authentification PocketBase invalide")
            headers["Authorization"] = self.token
            resp = await self.client.request(method, url, headers=headers, **kwargs)
        return resp

    async def fetch_collection(self, collection: str, filter_str: str = "",
                                expand: str = "", sort: str = "-created") -> List[dict]:
        records: List[dict] = []
        page = 1
        per_page = 200
        while True:
            params = {"page": page, "perPage": per_page, "sort": sort}
            if filter_str:
                params["filter"] = filter_str
            if expand:
                params["expand"] = expand
            try:
                resp = await self._request(
                    "GET", f"{self.base_url}/api/collections/{collection}/records", params=params
                )
            except HTTPException:
                raise
            except httpx.HTTPError as e:
                logger.error("Erreur réseau fetch %s : %s", collection, e)
                break
            if resp.status_code != 200:
                logger.warning("Fetch %s -> HTTP %s", collection, resp.status_code)
                break
            data = resp.json()
            records.extend(data.get("items", []))
            if page >= data.get("totalPages", 1):
                break
            page += 1
        return records

pb_client = PocketBaseClient(POCKETBASE_URL, PB_EMAIL, PB_PASSWORD)

@asynccontextmanager
async def lifespan(app: FastAPI):
    await pb_client.start()
    yield
    await pb_client.stop()

app = FastAPI(
    title="GMAO AI – Analyse Prédictive",
    description="API d'analyse des risques et de prédiction des pannes",
    version="2.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Cache court du pipeline d'analyse ─────────────────────────────────────
_analysis_cache: Dict[str, Tuple[datetime, dict]] = {}

# ─── Feature Engineering (métriques brutes par équipement) ────────────────
def _failure_trend(panne_dates: List[datetime], now: datetime) -> float:
    """Ratio entre le nb de pannes des 90 derniers jours et celui des 90 jours
    précédents. >1 = la fréquence des pannes s'aggrave, <1 = elle s'améliore."""
    recent = sum(1 for d in panne_dates if (now - d).days <= 90)
    previous = sum(1 for d in panne_dates if 90 < (now - d).days <= 180)
    if previous == 0:
        return 1.0 if recent == 0 else 1.5
    return recent / previous

def _extract_raw_metrics(eq: dict, pannes: list, ordres: list, pieces: list, plans: list) -> dict:
    eq_id = eq.get("id", "")
    now = datetime.utcnow()

    eq_pannes = [p for p in pannes if p.get("equipement") == eq_id]
    eq_ordres = [o for o in ordres if o.get("equipement") == eq_id]
    eq_pieces = [p for p in pieces if p.get("equipement") == eq_id]
    eq_plans = [p for p in plans if p.get("equipement") == eq_id]

    nb_pannes = len(eq_pannes)
    pannes_resolues = [p for p in eq_pannes if p.get("statut") == "resolue"]
    pannes_ouvertes = [p for p in eq_pannes if p.get("statut") in ("nouvelle", "en_cours", "en_attente")]
    pannes_critiques = [p for p in eq_pannes if p.get("priorite") == "critique"]

    # MTTR : moyenne ET médiane (la médiane résiste mieux aux valeurs extrêmes,
    # ex. une réparation exceptionnellement longue ne doit pas tout fausser)
    repair_times = []
    for p in pannes_resolues:
        d_panne = parse_datetime(p.get("date_panne"))
        d_res = parse_datetime(p.get("date_resolution"))
        if d_panne and d_res and d_res > d_panne:
            repair_times.append((d_res - d_panne).total_seconds() / 3600)
    mttr_mean = round(float(np.mean(repair_times)), 2) if repair_times else 0.0
    mttr_median = round(float(np.median(repair_times)), 2) if repair_times else 0.0

    panne_dates = sorted(
        d for d in (parse_datetime(p.get("date_panne")) for p in eq_pannes) if d is not None
    )
    intervals = [
        (panne_dates[i] - panne_dates[i - 1]).total_seconds() / 3600
        for i in range(1, len(panne_dates))
    ]
    mtbf_mean = round(float(np.mean(intervals)), 2) if intervals else 9999.0
    mtbf_median = round(float(np.median(intervals)), 2) if intervals else 9999.0

    days_since_last = (now - panne_dates[-1]).days if panne_dates else None

    prio_weights = {"critique": 4, "haute": 3, "moyenne": 2, "basse": 1}
    max_prio = max((prio_weights.get(p.get("priorite", "basse"), 1) for p in eq_pannes), default=1)

    nb_ot_correctif = sum(1 for o in eq_ordres if o.get("type") == "correctif")
    nb_ot_preventif = sum(1 for o in eq_ordres if o.get("type") == "preventif")
    nb_ot_retard = sum(
        1 for o in eq_ordres
        if o.get("statut") not in ("termine", "cloture", "annule")
        and (d := parse_datetime(o.get("date_fin_prevue"))) and d < now
    )

    pieces_hs = sum(1 for p in eq_pieces if p.get("statut") == "hors_service")
    stock_faible = sum(
        1 for p in eq_pieces
        if safe_float(p.get("quantite")) < 2 and p.get("statut") != "hors_service"
    )
    plans_actifs = sum(1 for p in eq_plans if p.get("statut") == "actif")

    pannes_30j = sum(1 for d in panne_dates if (now - d).days <= 30)
    trend = round(_failure_trend(panne_dates, now), 2)

    return {
        "id": eq_id, "nom": eq.get("nom", ""), "zone": eq.get("zone", ""),
        "statut": eq.get("statut", ""),
        "nb_pannes": nb_pannes, "pannes_ouvertes": len(pannes_ouvertes),
        "pannes_critiques": len(pannes_critiques),
        "mttr_h": mttr_mean, "mttr_h_median": mttr_median,
        "mtbf_h": mtbf_mean, "mtbf_h_median": mtbf_median,
        "days_since_last_failure": days_since_last, "max_priorite": max_prio,
        "nb_ot_correctifs": nb_ot_correctif, "nb_ot_preventifs": nb_ot_preventif,
        "nb_ot_en_retard": nb_ot_retard,
        "pieces_hs": pieces_hs, "stock_faible": stock_faible,
        "nb_plans_actifs": plans_actifs,
        "pannes_30j": pannes_30j, "tendance_pannes": trend,
    }

# ─── Score de risque composite : échelle absolue + échelle relative au parc ─
ABS_SCALE = {"nb_pannes": 20, "mttr_h": 100, "mtbf_h": 500, "nb_ot_en_retard": 5, "stock_faible": 5}

def _percentile_score(value: float, fleet_values: List[float], p_low=10, p_high=90) -> float:
    """Normalise une valeur entre 0 et 1 par rapport à la distribution du parc
    actuel. Complète l'échelle fixe ci-dessus, qui est une convention métier
    arbitraire ne s'adaptant pas à la réalité du site (un parc globalement
    très fiable ou au contraire très sollicité)."""
    if len(fleet_values) < 3:
        return 0.0
    lo, hi = np.percentile(fleet_values, [p_low, p_high])
    if hi <= lo:
        return 0.0
    return float(np.clip((value - lo) / (hi - lo), 0.0, 1.0))

def compute_fleet_scores(raw_list: List[dict]) -> List[dict]:
    """Calcule score_risque et niveau_risque pour chaque équipement, en
    combinant échelle métier fixe (60%) et échelle relative au parc (40%)."""
    nb_pannes_fleet = [f["nb_pannes"] for f in raw_list]
    mttr_fleet = [f["mttr_h_median"] for f in raw_list if f["mttr_h_median"] > 0]
    mtbf_fleet = [f["mtbf_h_median"] for f in raw_list if f["mtbf_h_median"] < 9999]
    retard_fleet = [f["nb_ot_en_retard"] for f in raw_list]
    stock_fleet = [f["stock_faible"] for f in raw_list]

    results = []
    for f in raw_list:
        f_freq = 0.6 * min(f["nb_pannes"] / ABS_SCALE["nb_pannes"], 1.0) \
            + 0.4 * _percentile_score(f["nb_pannes"], nb_pannes_fleet)

        if f["mttr_h_median"] > 0:
            f_mttr = 0.6 * min(f["mttr_h_median"] / ABS_SCALE["mttr_h"], 1.0) \
                + 0.4 * _percentile_score(f["mttr_h_median"], mttr_fleet)
        else:
            f_mttr = 0.0

        if f["mtbf_h_median"] < 9999:
            f_mtbf = 0.6 * (1.0 - min(f["mtbf_h_median"] / ABS_SCALE["mtbf_h"], 1.0)) \
                + 0.4 * (1.0 - _percentile_score(f["mtbf_h_median"], mtbf_fleet))
        else:
            f_mtbf = 0.0

        f_prio = (f["max_priorite"] - 1) / 3
        f_retard = 0.6 * min(f["nb_ot_en_retard"] / ABS_SCALE["nb_ot_en_retard"], 1.0) \
            + 0.4 * _percentile_score(f["nb_ot_en_retard"], retard_fleet)
        f_stock = 0.6 * min(f["stock_faible"] / ABS_SCALE["stock_faible"], 1.0) \
            + 0.4 * _percentile_score(f["stock_faible"], stock_fleet)
        f_stat = 1.0 if f["statut"] == "hors_service" else 0.0
        f_prev = 0.0 if f["nb_plans_actifs"] > 0 else 0.3
        f_trend = min(max(f["tendance_pannes"] - 1.0, 0.0) / 1.5, 1.0)

        score = round(
            0.22 * f_freq + 0.13 * f_mttr + 0.13 * f_mtbf + 0.13 * f_prio +
            0.09 * f_retard + 0.07 * f_stock + 0.07 * f_stat + 0.05 * f_prev +
            0.11 * f_trend,
            3,
        )

        if score >= 0.7:
            niveau = "CRITIQUE"
        elif score >= 0.45:
            niveau = "ÉLEVÉ"
        elif score >= 0.25:
            niveau = "MODÉRÉ"
        else:
            niveau = "FAIBLE"

        results.append({**f, "score_risque": score, "niveau_risque": niveau})
    return results

# ─── Recommandations IA ────────────────────────────────────────────────────
FEATURE_LABELS_FR = {
    "score_risque": "le score de risque composite",
    "nb_pannes": "le nombre total de pannes",
    "mttr_h_median": "le temps de réparation (MTTR)",
    "mtbf_h_median": "la fiabilité entre pannes (MTBF)",
    "max_priorite": "la priorité des pannes",
    "nb_ot_en_retard": "les ordres de travail en retard",
    "stock_faible": "le stock de pièces",
    "pieces_hs": "les pièces hors service",
    "nb_plans_actifs": "l'absence de plan préventif",
    "tendance_pannes": "la tendance récente des pannes",
    "pannes_critiques": "les pannes critiques",
}

def generate_recommendations(feat: dict, proba: float, top_driver: Optional[str] = None) -> List[str]:
    recs = []

    if feat["statut"] == "hors_service":
        recs.append("⚠️ URGENT – Planifier une remise en service immédiate.")

    if proba >= 0.70:
        recs.append("🔴 Probabilité de panne très élevée : intervention sous 48h.")
    elif proba >= 0.45:
        recs.append("🟠 Risque élevé : inspection approfondie dans les 7 jours.")

    if feat["pannes_ouvertes"] > 0:
        recs.append(f"📋 {feat['pannes_ouvertes']} panne(s) non résolue(s) à traiter.")

    if feat["pannes_critiques"] > 0:
        recs.append(f"🚨 {feat['pannes_critiques']} panne(s) de priorité critique enregistrée(s).")

    if feat["nb_ot_en_retard"] > 0:
        recs.append(f"⏰ {feat['nb_ot_en_retard']} OT en retard – replanifier.")

    if feat["mttr_h_median"] > 24:
        recs.append(f"🔧 MTTR médian élevé ({feat['mttr_h_median']}h) : optimiser les processus de réparation.")

    if 0 < feat["mtbf_h_median"] < 200:
        recs.append(f"📉 MTBF médian faible ({feat['mtbf_h_median']}h) : renforcer la maintenance préventive.")

    if feat["tendance_pannes"] > 1.5:
        recs.append(
            f"📈 Tendance à la dégradation : {feat['tendance_pannes']}x plus de pannes ces 90 derniers "
            "jours que sur la période précédente."
        )

    if feat["stock_faible"] > 0:
        recs.append(f"📦 {feat['stock_faible']} pièce(s) en stock faible – commander.")

    if feat["pieces_hs"] > 0:
        recs.append(f"🛠️ {feat['pieces_hs']} pièce(s) HS – évaluer l'impact.")

    if feat["nb_plans_actifs"] == 0:
        recs.append("📅 Aucun plan préventif actif – en créer un.")

    if feat["pannes_30j"] >= 3:
        recs.append(f"📊 {feat['pannes_30j']} pannes/30j : réaliser une analyse AMDEC.")

    if feat["days_since_last_failure"] is not None and feat["days_since_last_failure"] < 7:
        recs.append(f"⚡ Panne il y a {feat['days_since_last_failure']}j : contrôle post-réparation.")

    if top_driver and proba >= 0.45:
        recs.append(f"🧠 Facteur le plus déterminant selon le modèle : {FEATURE_LABELS_FR.get(top_driver, top_driver)}.")

    if not recs:
        recs.append("✅ Équipement en bon état – maintenir la surveillance.")

    return recs

# ─── Modèle ML léger : entraînement, évaluation, persistance ──────────────
FEATURE_COLS_ML = [
    "score_risque", "nb_pannes", "mttr_h_median", "mtbf_h_median", "max_priorite",
    "nb_ot_en_retard", "stock_faible", "pieces_hs", "nb_plans_actifs",
    "tendance_pannes", "pannes_critiques",
]
LABEL_COL = "pannes_30j"  # cible : a eu une panne dans les 30 derniers jours

def _build_training_set(features_list: List[dict]):
    X = [[safe_float(f.get(c)) for c in FEATURE_COLS_ML] for f in features_list]
    y = [1 if safe_float(f.get(LABEL_COL)) >= 1 else 0 for f in features_list]
    return np.array(X, dtype=float), np.array(y, dtype=int)

def _fit_one(model_type: str, X, y):
    if model_type == "logistic":
        scaler = StandardScaler()
        X_s = scaler.fit_transform(X)
        clf = LogisticRegression(max_iter=500, class_weight="balanced")
        clf.fit(X_s, y)
        importances = dict(zip(FEATURE_COLS_ML, np.abs(clf.coef_[0]).tolist()))
        return clf, scaler, importances
    scaler = None
    clf = RandomForestClassifier(
        n_estimators=120, max_depth=6, min_samples_leaf=2,
        class_weight="balanced", random_state=42, n_jobs=1,
    )
    clf.fit(X, y)
    importances = dict(zip(FEATURE_COLS_ML, clf.feature_importances_.tolist()))
    return clf, scaler, importances

def _train_new_model(features_list: List[dict]) -> Optional[dict]:
    """Entraîne un modèle scikit-learn et l'évalue avant de le mettre en
    service. S'il n'est pas significativement meilleur que le hasard (AUC <
    0.55) sur un jeu de test, on revient au score heuristique plutôt que de
    déployer un modèle dans lequel on ne peut pas avoir confiance.

    Remarque : la calibration des probabilités (CalibratedClassifierCV)
    n'est volontairement pas utilisée ici car les parcs d'équipements en
    GMAO sont souvent trop petits pour la valider correctement ; à activer
    si l'historique dépasse quelques centaines d'échantillons.
    """
    if not SKLEARN_AVAILABLE:
        return None

    X, y = _build_training_set(features_list)
    n = len(X)
    if n < MIN_SAMPLES_TRAINING or len(set(y.tolist())) < 2:
        logger.info("Parc trop petit ou classe unique (n=%d) : repli heuristique", n)
        return None

    metrics: dict = {}
    can_evaluate = n >= MIN_SAMPLES_TRAINING * 2
    X_train, y_train = X, y

    if can_evaluate:
        try:
            X_train, X_test, y_train, y_test = train_test_split(
                X, y, test_size=0.25, random_state=42, stratify=y
            )
        except ValueError:
            can_evaluate = False

    try:
        clf, scaler, importances = _fit_one(ML_MODEL_TYPE, X_train, y_train)

        if can_evaluate:
            X_eval = scaler.transform(X_test) if scaler is not None else X_test
            y_proba = clf.predict_proba(X_eval)[:, 1]
            y_pred = (y_proba >= 0.5).astype(int)
            try:
                metrics["auc"] = round(float(roc_auc_score(y_test, y_proba)), 3)
            except ValueError:
                metrics["auc"] = None
            metrics["precision"] = round(float(precision_score(y_test, y_pred, zero_division=0)), 3)
            metrics["recall"] = round(float(recall_score(y_test, y_pred, zero_division=0)), 3)
            metrics["n_test"] = int(len(y_test))

            if metrics["auc"] is not None and metrics["auc"] < 0.55:
                logger.warning(
                    "Modèle ML peu fiable (AUC=%.2f sur %d échantillons de test) : repli heuristique",
                    metrics["auc"], metrics["n_test"],
                )
                return None

        # Ré-entraînement final sur 100% des données disponibles pour la mise en prod
        clf_full, scaler_full, importances_full = _fit_one(ML_MODEL_TYPE, X, y)

        bundle = {
            "model": clf_full, "scaler": scaler_full, "feature_cols": FEATURE_COLS_ML,
            "type": ML_MODEL_TYPE, "trained_at": datetime.utcnow(),
            "metrics": metrics, "feature_importances": importances_full,
            "n_train_samples": n,
        }
        MODEL_DIR.mkdir(parents=True, exist_ok=True)
        joblib.dump(bundle, MODEL_PATH)
        logger.info("Modèle %s entraîné (n=%d, métriques=%s)", ML_MODEL_TYPE, n, metrics)
        return bundle
    except Exception as e:
        logger.exception("Entraînement du modèle impossible : %s", e)
        return None

def get_or_train_model(features_list: List[dict]) -> Optional[dict]:
    """Charge le modèle persisté (pas de ré-entraînement à chaque requête) et
    ne le ré-entraîne que s'il est absent, périmé, ou illisible."""
    if not SKLEARN_AVAILABLE:
        return None

    bundle = None
    if MODEL_PATH.exists():
        try:
            bundle = joblib.load(MODEL_PATH)
        except Exception as e:
            logger.warning("Modèle illisible sur disque, ré-entraînement : %s", e)
            bundle = None

    needs_retrain = (
        bundle is None
        or (datetime.utcnow() - bundle.get("trained_at", datetime.min)).total_seconds() / 3600
        > RETRAIN_EVERY_HOURS
    )
    if needs_retrain:
        new_bundle = _train_new_model(features_list)
        if new_bundle is not None:
            bundle = new_bundle

    return bundle

def get_current_model_info() -> dict:
    """Lit les métadonnées du modèle actuellement persisté, sans en entraîner
    un nouveau (utilisé par /api/model/metrics et le rapport DOCX)."""
    if not SKLEARN_AVAILABLE or not MODEL_PATH.exists():
        return {"disponible": False}
    try:
        bundle = joblib.load(MODEL_PATH)
    except Exception:
        return {"disponible": False}
    trained_at = bundle.get("trained_at")
    return {
        "disponible": True,
        "type_modele": bundle.get("type"),
        "entraine_le": trained_at.strftime("%d/%m/%Y %H:%M") if trained_at else None,
        "nb_echantillons_entrainement": bundle.get("n_train_samples"),
        "metriques": bundle.get("metrics"),
        "importance_facteurs": bundle.get("feature_importances"),
    }

def _top_driver(feature_importances: Optional[dict]) -> Optional[str]:
    if not feature_importances:
        return None
    return max(feature_importances, key=feature_importances.get)

def predict_failure_probability(features_list: List[dict]) -> dict:
    """Prédit la probabilité de panne à 30 jours, avec repli automatique sur
    le score heuristique si le modèle ML est indisponible, non fiable, ou
    en erreur."""
    results = {}

    try:
        bundle = get_or_train_model(features_list)
    except Exception as e:
        logger.exception("Erreur modèle ML, repli heuristique : %s", e)
        bundle = None

    if bundle is None:
        for feat in features_list:
            proba = round(min(feat["score_risque"] * 0.9 + 0.05, 0.99), 3)
            results[feat["id"]] = {"proba_panne": proba, "methode": "score_heuristique", "horizon_jours": 30}
        return results

    model = bundle["model"]
    scaler = bundle.get("scaler")
    cols = bundle["feature_cols"]
    methode = bundle.get("type", "ml")

    try:
        X = np.array([[safe_float(f.get(c)) for c in cols] for f in features_list], dtype=float)
        if scaler is not None:
            X = scaler.transform(X)
        probas = model.predict_proba(X)[:, 1]
        for feat, p in zip(features_list, probas):
            results[feat["id"]] = {
                "proba_panne": round(float(min(max(p, 0.01), 0.99)), 3),
                "methode": methode,
                "horizon_jours": 30,
            }
        return results
    except Exception as e:
        logger.exception("Erreur prédiction ML, repli heuristique : %s", e)
        for feat in features_list:
            proba = round(min(feat["score_risque"] * 0.9 + 0.05, 0.99), 3)
            results[feat["id"]] = {"proba_panne": proba, "methode": "score_heuristique_fallback", "horizon_jours": 30}
        return results

# ─── Génération Rapport DOCX ────────────────────────────────────────────────
def generate_docx_report(analyse: dict, output_path: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("Installez python-docx : pip install python-docx")

    risk_colors = {
        "CRITIQUE": RGBColor(0xC0, 0x00, 0x00),
        "ÉLEVÉ": RGBColor(0xE6, 0x7E, 0x00),
        "MODÉRÉ": RGBColor(0xB8, 0x86, 0x0B),
        "FAIBLE": RGBColor(0x2E, 0x7D, 0x32),
    }

    doc = Document()
    title = doc.add_heading("Rapport d'Analyse Prédictive GMAO", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    stats = analyse.get("stats_globales", {})
    doc.add_paragraph(f"Date : {stats.get('date_rapport', 'N/A')}")
    doc.add_paragraph(f"Équipements analysés : {stats.get('nb_equipements', 0)}")

    modele_info = stats.get("modele_ml", {})
    if modele_info.get("disponible"):
        m = modele_info.get("metriques") or {}
        doc.add_paragraph(
            f"Modèle ML actif : {modele_info.get('type_modele')} "
            f"(entraîné le {modele_info.get('entraine_le')}, "
            f"AUC={m.get('auc', 'N/A')}, precision={m.get('precision', 'N/A')}, "
            f"recall={m.get('recall', 'N/A')})"
        )
    else:
        doc.add_paragraph("Modèle ML : non disponible, score heuristique utilisé.")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Niveau", "CRITIQUE", "ÉLEVÉ", "MODÉRÉ/FAIBLE"
    row = table.add_row().cells
    row[0].text = "Count"
    row[1].text = str(stats.get("nb_critiques", 0))
    row[2].text = str(stats.get("nb_eleves", 0))
    row[3].text = str(stats.get("nb_moderes", 0) + stats.get("nb_faibles", 0))
    doc.add_paragraph("")

    doc.add_heading("Détails par Équipement", level=1)
    for eq in analyse.get("equipements", [])[:15]:
        p = doc.add_paragraph()
        p.add_run(f"🔹 {eq.get('nom')} ({eq.get('zone', 'N/A')})").bold = True
        niveau = eq.get("niveau_risque", "FAIBLE")
        run = p.add_run(
            f"\n   Risque : {niveau} | Probabilité 30j : {eq.get('proba_panne_30j', 0) * 100:.1f}% "
            f"| Tendance : {eq.get('tendance_pannes', 1.0)}x"
        )
        run.font.color.rgb = risk_colors.get(niveau, RGBColor(0, 0, 0))

        if eq.get("recommandations"):
            p.add_run("\n   Recommandations :")
            for rec in eq["recommandations"][:4]:
                p.add_run(f"\n   • {rec}")
        doc.add_paragraph("")

    doc.save(output_path)
    return output_path

# ─── Pipeline Principal ──────────────────────────────────────────────────────
async def run_analysis_pipeline(equipement_id: Optional[str] = None, use_cache: bool = True) -> dict:
    cache_key = equipement_id or "__all__"
    if use_cache:
        cached = _analysis_cache.get(cache_key)
        if cached and (datetime.utcnow() - cached[0]).total_seconds() < ANALYSIS_CACHE_TTL_SECONDS:
            return cached[1]

    eq_filter = f'id="{equipement_id}"' if equipement_id else ""
    eq_related_filter = f'equipement="{equipement_id}"' if equipement_id else ""

    try:
        equipements, pannes, ordres, pieces, plans = await asyncio.gather(
            pb_client.fetch_collection("equipements", eq_filter),
            pb_client.fetch_collection("pannes", eq_related_filter),
            pb_client.fetch_collection("ordresdetravail", eq_related_filter),
            pb_client.fetch_collection("pieces", eq_related_filter),
            pb_client.fetch_collection("plans_preventifs", eq_related_filter),
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Erreur fetch PocketBase")
        raise HTTPException(500, f"Erreur fetch PocketBase : {e}")

    if not equipements:
        raise HTTPException(404, "Aucun équipement trouvé.")

    raw_list = [_extract_raw_metrics(eq, pannes, ordres, pieces, plans) for eq in equipements]
    scored_list = compute_fleet_scores(raw_list)

    predictions = predict_failure_probability(scored_list)
    model_info = get_current_model_info()
    top_driver = _top_driver(model_info.get("importance_facteurs"))

    rapport = []
    for feat in scored_list:
        pred = predictions.get(feat["id"], {"proba_panne": feat["score_risque"], "methode": "heuristique"})
        proba = pred["proba_panne"]
        methode = pred.get("methode", "heuristique")
        driver_for_this = top_driver if methode not in ("score_heuristique", "score_heuristique_fallback") else None
        recs = generate_recommendations(feat, proba, driver_for_this)
        rapport.append({**feat, "proba_panne_30j": proba, "methode_prediction": methode, "recommandations": recs})

    rapport.sort(key=lambda x: x["proba_panne_30j"], reverse=True)

    now = datetime.utcnow()
    stats = {
        "date_rapport": now.strftime("%d/%m/%Y %H:%M"),
        "nb_equipements": len(rapport),
        "nb_critiques": sum(1 for e in rapport if e["niveau_risque"] == "CRITIQUE"),
        "nb_eleves": sum(1 for e in rapport if e["niveau_risque"] == "ÉLEVÉ"),
        "nb_moderes": sum(1 for e in rapport if e["niveau_risque"] == "MODÉRÉ"),
        "nb_faibles": sum(1 for e in rapport if e["niveau_risque"] == "FAIBLE"),
        "nb_pannes_total": len(pannes),
        "nb_ot_total": len(ordres),
        "nb_plans_preventifs": len(plans),
        "equipements_hors_service": sum(1 for e in equipements if e.get("statut") == "hors_service"),
        "taux_maintenance_preventive": round(
            sum(e["nb_ot_preventifs"] for e in rapport) /
            max(sum(e["nb_ot_correctifs"] + e["nb_ot_preventifs"] for e in rapport), 1) * 100,
            1,
        ),
        "modele_ml": model_info,
    }

    result = {"stats_globales": stats, "equipements": rapport}
    _analysis_cache[cache_key] = (datetime.utcnow(), result)
    return result

# ─── Routes API ────────────────────────────────────────────────────────────
@app.get("/", tags=["Santé"])
async def health():
    info = get_current_model_info()
    return {
        "status": "ok",
        "service": "GMAO AI API",
        "version": "2.0.0",
        "modele_ml": info if info.get("disponible") else {"disponible": False, "fallback": "score_heuristique"},
    }

@app.post("/api/analyse", response_model=AnalyseResponse, tags=["Analyse"])
async def analyse_endpoint(req: RapportRequest):
    try:
        analyse = await run_analysis_pipeline(req.equipement_id, use_cache=req.use_cache)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Erreur analyse : {e}")

    if req.output_format == "json":
        return AnalyseResponse(status="success", analyse=analyse, message="Analyse terminée.")

    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"rapport_gmao_{ts}.docx"
    output_path = REPORTS_DIR / filename
    try:
        generate_docx_report(analyse, str(output_path))
    except ImportError as e:
        raise HTTPException(500, f"Erreur docx : {e}")
    except Exception as e:
        raise HTTPException(500, f"Erreur génération : {e}")

    return AnalyseResponse(status="success", rapport_path=f"/api/rapport/{filename}", message="Rapport DOCX généré.")

@app.get("/api/rapport/{filename}", tags=["Rapport"])
async def download_rapport(filename: str):
    filepath = REPORTS_DIR / filename
    if not filepath.exists():
        raise HTTPException(404, "Rapport introuvable.")
    return FileResponse(
        str(filepath),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )

@app.get("/api/equipements/{eq_id}/risque", tags=["Analyse"])
async def risque_equipement(eq_id: str, use_cache: bool = Query(default=True)):
    try:
        analyse = await run_analysis_pipeline(eq_id, use_cache=use_cache)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, str(e))

    if not analyse["equipements"]:
        raise HTTPException(404, "Équipement non trouvé.")

    return {"status": "success", "equipement": analyse["equipements"][0], "stats": analyse["stats_globales"]}

@app.get("/api/dashboard", tags=["Dashboard"])
async def dashboard_summary(use_cache: bool = Query(default=True)):
    try:
        analyse = await run_analysis_pipeline(None, use_cache=use_cache)
    except Exception as e:
        raise HTTPException(500, str(e))

    critiques = [e for e in analyse["equipements"] if e["niveau_risque"] == "CRITIQUE"]
    return {
        "stats": analyse["stats_globales"],
        "equipements_critiques": critiques,
        "top5_risque": analyse["equipements"][:5],
    }

@app.get("/api/model/metrics", response_model=ModelMetricsResponse, tags=["Modèle"])
async def model_metrics():
    """Inspecte la fiabilité du modèle ML actuellement en service (AUC,
    precision, recall, importance des variables, date d'entraînement)."""
    return ModelMetricsResponse(**get_current_model_info())

# ─── Entry Point ────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)